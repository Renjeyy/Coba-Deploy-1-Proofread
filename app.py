import os
import io
import re
import json
import zipfile
import difflib
import fitz
import docx
import pandas as pd
import google.generativeai as genai
import shutil 
from flask import (
    Flask, request, jsonify, render_template, send_file, 
    make_response, redirect, url_for, flash
)
from dotenv import load_dotenv
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, Inches
from docx import Document 

from flask_sqlalchemy import SQLAlchemy
from flask_login import (
    LoginManager, UserMixin, login_user, logout_user, login_required, current_user
)
from werkzeug.security import generate_password_hash, check_password_hash
import datetime 
load_dotenv()

app = Flask(__name__)

app.config['SECRET_KEY'] = 'kunci-rahasia-anda-yang-acak-dan-kuat'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

APP_ROOT = os.path.dirname(os.path.abspath(__file__))
app.config['UPLOAD_FOLDER'] = os.path.join(APP_ROOT, 'data') 

db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login' 
login_manager.login_message = 'Silakan login untuk mengakses halaman ini.'

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

try:
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise ValueError("GOOGLE_API_KEY tidak ditemukan di file .env")
    
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.5-pro') 
except Exception as e:
    print(f"Error saat mengkonfigurasi Google AI: {e}")

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    label = db.Column(db.String(100), nullable=False)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)
    
    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class AnalysisLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    
    # Relasi ke user
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    
    # Data Analisis
    filename = db.Column(db.String(255), nullable=False)
    feature_type = db.Column(db.String(50), nullable=False)
    document_type = db.Column(db.String(100), nullable=True)
    
    # Waktu dan Status
    start_time = db.Column(db.DateTime, nullable=False, default=datetime.datetime.utcnow)
    end_time = db.Column(db.DateTime, nullable=True)
    deadline = db.Column(db.DateTime, nullable=True) # <<<< TAMBAHKAN INI
    status = db.Column(db.String(20), nullable=False, default='unfinished') # unfinished, done, error, manual, overdue

    # Relasi
    user = db.relationship('User', backref='analysis_logs')

# --- PERBAIKAN TOTAL: Menambahkan kolom folder_type ---
class SharedFolder(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    folder_name = db.Column(db.String(200), nullable=False)
    owner_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    shared_with_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    # FIX: Tambahkan kolom yang dibutuhkan database Anda
    folder_type = db.Column(db.String(50), nullable=False, default='general') 
    
    # Tambahkan relasi untuk memudahkan query
    owner = db.relationship('User', foreign_keys=[owner_id], backref='owned_shares')
    shared_with_user = db.relationship('User', foreign_keys=[shared_with_id], backref='received_shares')
# --- AKHIR PERBAIKAN TOTAL ---

# >>>>>> REVISION START: Model Comment yang diperbarui untuk threaded comments <<<<<<
class Comment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    
    # Kunci untuk mengidentifikasi hasil analisis yang dikomentari
    owner_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False) # ID pemilik hasil
    folder_name = db.Column(db.String(200), nullable=False)
    filename = db.Column(db.String(255), nullable=False)
    row_id = db.Column(db.Integer, nullable=False) # ID baris dalam tabel hasil

    # Data Komentar
    username = db.Column(db.String(100), nullable=False) # Nama user yang memberi komentar
    text = db.Column(db.Text, nullable=False)
    timestamp = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    
    # >>>>>> TAMBAHKAN BARIS INI <<<<<<
    # Untuk menyimpan komentar mana yang dibalas (jika ini adalah reply)
    parent_id = db.Column(db.Integer, db.ForeignKey('comment.id'), nullable=True)
    
    # Relasi untuk memudahkan query (opsional tapi sangat membantu)
    replies = db.relationship('Comment', backref=db.backref('parent', remote_side=[id]), lazy='dynamic')
    
    __table_args__ = (db.UniqueConstraint('owner_id', 'folder_name', 'filename', 'row_id', name='_unique_comment_target'),)

class RowAction(db.Model):
    id = db.Column(db.Integer, primary_key=True)

    owner_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    folder_name = db.Column(db.String(200), nullable=False)
    filename = db.Column(db.String(255), nullable=False)

    row_id = db.Column(db.Integer, nullable=False)
    
    # Data yang disimpan
    is_ganti = db.Column(db.Boolean, default=False, nullable=False)
    pic_user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    
    # Relasi
    pic_user = db.relationship('User', foreign_keys=[pic_user_id])
    
    __table_args__ = (db.UniqueConstraint('owner_id', 'folder_name', 'filename', 'row_id', name='_unique_row_action'),)

class Message(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    sender_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    recipient_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    subject = db.Column(db.String(255), nullable=False)
    body = db.Column(db.Text, nullable=True)
    attachment_path = db.Column(db.String(500), nullable=True) # Path di server
    original_filename = db.Column(db.String(255), nullable=True) # Nama asli file
    timestamp = db.Column(db.DateTime, nullable=False, default=datetime.datetime.utcnow)
    is_read = db.Column(db.Boolean, default=False, nullable=False)

    # Relasi
    sender = db.relationship('User', foreign_keys=[sender_id], backref='sent_messages')
    recipient = db.relationship('User', foreign_keys=[recipient_id], backref='received_messages')

# >>>>>> REVISION END <<<<<<

@login_manager.user_loader
def load_user(user_id):
    """Fungsi wajib untuk Flask-Login"""
    return User.query.get(int(user_id))

@app.cli.command("create-db")
def create_db_command():
    """Membuat database dan 5 user awal."""
    with app.app_context():
        db.create_all() 
        
        users_data = [
            ("Deny Syahbani", "Deny1234", "CoE"),
            ("Jihan Abigail", "Jihan1234", "CoE"),
            ("Fadian Dwiantara", "Fadian1234", "CoE"),
            ("Winda Anggraeni", "Winda1234", "CoE"),
            ("Renzie Aditya", "Renzie1234", "CoE")
        ]
        
        for username, password, label in users_data:
            if not User.query.filter_by(username=username).first():
                user = User(username=username, label=label)
                user.set_password(password)
                db.session.add(user)
                print(f"User {username} dibuat.")
        
        db.session.commit()
        print("Database dan user awal telah selesai dibuat.")
# ==============================================================================


# ==============================================================================
# ===         FUNGSI HELPER BARU: MANAJEMEN FOLDER        ===
# ==============================================================================

def get_user_root_folder():
    """Mendapatkan path folder root pengguna saat ini."""
    if not current_user.is_authenticated:
        return None
    
    # Pastikan ID pengguna aman untuk dijadikan nama folder
    user_id_str = str(current_user.id)
    if not user_id_str.isalnum():
        raise ValueError("User ID tidak valid untuk path folder.")
        
    user_folder_path = os.path.join(app.config['UPLOAD_FOLDER'], user_id_str)
    
    # Pastikan folder user ini ada SETIAP KALI mereka mengakses
    os.makedirs(user_folder_path, exist_ok=True)
    return user_folder_path

def create_user_folder(folder_name):
    """Membuat sub-folder di folder root pengguna."""
    root_folder = get_user_root_folder()
    if not root_folder:
        raise Exception("Pengguna tidak terautentikasi.")
    
    # Membersihkan nama folder untuk keamanan sistem file
    clean_folder_name = re.sub(r'[^\w\s-]', '', folder_name).strip() # Hapus karakter aneh
    clean_folder_name = re.sub(r'[-\s]+', '_', clean_folder_name) # Ganti spasi dengan underscore
    
    if not clean_folder_name:
        raise ValueError("Nama folder tidak valid setelah dibersihkan.")

    new_folder_path = os.path.join(root_folder, clean_folder_name)
    if os.path.exists(new_folder_path):
        raise ValueError("Folder dengan nama yang sama sudah ada.")

    os.makedirs(new_folder_path)
    return clean_folder_name

def get_user_folders():
    """
    MODIFIKASI: Mengambil folder milik sendiri DAN folder yang di-share.
    Mengembalikan list of dictionaries.
    """
    root_folder = get_user_root_folder()
    if not root_folder:
        return []
    
    my_folders = []
    
    # 1. Ambil folder milik sendiri (dari sistem file)
    try:
        owned_folders = [
            d for d in os.listdir(root_folder) 
            if os.path.isdir(os.path.join(root_folder, d))
        ]
        for folder in owned_folders:
            my_folders.append({
                "name": folder,
                "owner_name": current_user.username,
                "is_owner": True,
                "owner_id": current_user.id # ID kita sendiri
            })
    except Exception as e:
        print(f"Error saat listing folder di {root_folder}: {e}")

    # 2. Ambil folder yang di-share ke kita (dari DB)
    shared_folders = SharedFolder.query.filter_by(shared_with_id=current_user.id).all()
    
    for share in shared_folders:
        # Pastikan folder fisik-nya masih ada
        owner_root_path = os.path.join(app.config['UPLOAD_FOLDER'], str(share.owner_id))
        folder_path = os.path.join(owner_root_path, share.folder_name)
        
        if os.path.isdir(folder_path):
            my_folders.append({
                "name": share.folder_name,
                "owner_name": share.owner.username, # Menggunakan relasi
                "is_owner": False,
                "owner_id": share.owner_id # ID pemilik
            })

    return my_folders

def _extract_text_with_pages(file_bytes, file_extension):
    pages_content = []
    if file_extension == 'pdf':
        try:
            pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
            for page_num, page in enumerate(pdf_document):
                pages_content.append({"halaman": page_num + 1, "teks": page.get_text()})
            pdf_document.close()
        except Exception as e:
            raise ValueError(f"Gagal membaca file PDF: {e}")
    elif file_extension == 'docx':
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = "\n".join([para.text for para in doc.paragraphs])
            pages_content.append({"halaman": 1, "teks": full_text})
        except Exception as e:
            raise ValueError(f"Gagal membaca file DOCX: {e}")
    else:
        raise ValueError("Format file tidak didukung. Harap unggah .pdf atau .docx")
    return pages_content

def _get_text_from_flask_file(file):
    file_bytes = file.read()
    file.seek(0) 
    file_extension = file.filename.split('.')[-1].lower()
    return _extract_text_with_pages(file_bytes, file_extension)

def extract_sentences_with_pages(pages_content):
    sentences_with_pages = []
    for page in pages_content:
        page_num = page['halaman']
        text = page['teks']
        sentences = split_text_into_sentences(text) 
        for sentence in sentences:
            sentences_with_pages.append({'sentence': sentence, 'page': page_num})
    return sentences_with_pages

def _get_full_text_from_file(file):
    file_bytes = file.read()
    file.seek(0)
    file_extension = file.filename.split('.')[-1].lower()
    
    pages = _extract_text_with_pages(file_bytes, file_extension)
    return "\n".join([page['teks'] for page in pages])

def proofread_with_gemini(text_to_check):
    if not text_to_check or text_to_check.isspace():
        return []
    prompt = f"""
    Anda adalah seorang auditor dan ahli bahasa Indonesia yang sangat teliti. Anda diberikan dokumen dan tugas Anda adalah melakukan proofread pada teks berikut. Fokus pada:
    1. Memperbaiki kesalahan ketik (typo) agar semuanya sesuai dengan standar KBBI dan PUEBI.
    1. Kalau ada kata-kata yang tidak sesuai KBBI dan PUEBI, tolong jangan highlight semua kalimatnya, tapi cukup highlight kata-kata yang salah serta perbaiki kata-kata itu aja, jangan perbaiki semua kalimatnya
    3. Jika ada kata yang diitalic, biarkan saja
    4. Nama-nama yang diberi ini pastikan benar juga "Yullyan, I Made Suandi Putra, Laila Fajriani, Hari Sundoro, Bakhas Nasrani Diso, Rizky Ananda Putra, Wirawan Arief Nugroho, Lelya Novita Kusumawati, Ryani Ariesti Syafitri, Darmo Saputro Wibowo, Lucky Parwitasari, Handarudigdaya Jalanidhi Kuncaratrah, Fajar Setianto, Jaka Tirtana Hanafiah, tMuhammad Rosyid Ridho Muttaqien, Octovian Abrianto, Deny Sjahbani, Jihan Abigail, Winda Anggraini, Fadian Dwiantara, Aliya Anindhita Rachman"
    5. Fontnya arial dan jangan diganti. Khusus untuk judul paling atas, itu font sizenya 12 dan bodynya selalu 11
    6. Khusus "Indonesia Financial Group (IFG)", meskipun bahasa inggris, tidak perlu di italic
    7. Kalau ada kata yang sudah diberikan akronimnya di awal, maka di halaman berikut-berikutnya cukup akronimnya saja, tidak perlu ditulis lengkap lagi
    8. Pada bagian Nomor surat dan Penutup tidak perlu dicek, biarkan seperti itu
    9. Ketika Anda perbaiki, fontnya pastikan Arial dengan ukuran 11 juga (Tidak diganti)
    10. Pada kalimat "Indonesia Financial Group", jika terdapat kata typo "Finansial", tolong Anda sarankan untuk ganti ke "Financial"
    11. Yang benar adalah "Satuan Kerja Audit Internal", bukan "Satuan Pengendali Internal Audit"
    12. Jika terdapat kata "reviu", biarkan itu sebagai benar
    13. Kalau ada kata "IM", "ST", "SKAI", "IFG", "TV (Angka Romawi)", "RKAT", dan "RKAP", itu tidak perlu ditandai sebagai salah dan tidak perlu disarankan untuk italic / bold / underline
    14. Untuk nama modul seperti "Modul Sourcing, dll", itu tidak perlu italic
    15. Kalau ada kata dalam bahasa inggris yang masih masuk akal dan nyambung dengan kalimat yang dibahas, tidak perlu Anda sarankan untuk ganti ke bahasa indonesia
    16. Jika ada bahasa inggris dan akronimnya seperti "General Ledger (GL)", tolong dilakukan italic pada kata tersebut pada saat download file hasil revisinya, akronimnya tidak perlu diitalic
    17. Awal kalimat selalu dimulai dengan huruf kapital. Jika akhir poin diberi tanda ";", maka poin selanjutnya tidak perlu kapital
    18. Di file hasil revisi, Anda jangan ganti dari yang aslinya. Misalnya kalau ada kata yang diitalic di file asli, jangan Anda hilangkan italicnya
    19. Tolong perhatikan juga tanda bacanya, seperti koma, titik koma, titik, tanda hubung, dan lain-lain. Pastikan sesuai dan ada tanda titik di setiap akhir kalimat
    20. Biarkan kata "Advisory" sebagai "Advisory" saja, tidak perlu diganti
    21. Kalau ada bahasa inggris yang belum diitalic, tolong diitalic

    PENTING: Berikan hasil dalam format yang SANGAT KETAT. Untuk setiap kesalahan, gunakan format:
    [SALAH] kata atau frasa yang salah -> [BENAR] kata atau frasa perbaikan -> [KALIMAT] kalimat lengkap asli tempat kesalahan ditemukan

    Contoh:
    [SALAH] dikarenakan -> [BENAR] karena -> [KALIMAT] Hal itu terjadi dikarenakan kelalaian petugas.

    Jika tidak ada kesalahan sama sekali, kembalikan teks: "TIDAK ADA KESALAHAN"

    Berikut adalah teks yang harus Anda periksa:
    ---
    {text_to_check}
    """
    try:
        response = model.generate_content(prompt)
        pattern = re.compile(r"\[SALAH\]\s*(.*?)\s*->\s*\[BENAR\]\s*(.*?)\s*->\s*\[KALIMAT\]\s*(.*?)\s*(\n|$)", re.IGNORECASE | re.DOTALL)
        found_errors = pattern.findall(response.text)
        return [{"salah": salah.strip(), "benar": benar.strip(), "kalimat": kalimat.strip()} for salah, benar, kalimat, _ in found_errors]
    except Exception as e:
        print(f"Terjadi kesalahan saat menghubungi AI: {e}")
        return [{"salah": "ERROR", "benar": str(e), "kalimat": "Gagal menghubungi API"}]

def split_text_into_sentences(full_text):
    """
    Memecah teks penuh menjadi daftar kalimat.
    Ini adalah implementasi sederhana dan mungkin tidak sempurna untuk semua kasus,
    tetapi merupakan awal yang baik untuk dokumen formal.
    """
    if not full_text:
        return []

    sentences = re.split(r'(?<=[.!?])\s+', full_text)
    return [s.strip() for s in sentences if len(s.strip()) > 10]

def analyze_document_by_section(original_text, revised_text):
    """
    Menganalisis kesesuaian makna antara dokumen asli dan revisi berdasarkan sub-bab.
    Fungsi ini telah direvisi untuk meminta AI memberikan alasan dalam dua poin terpisah.
    """
    if not original_text or not revised_text:
        return []

    # PROMPT YANG TELAH DIREVISI
    prompt = f"""
    Anda adalah seorang auditor ahli. Tugas Anda adalah membandingkan dua dokumen: Dokumen Asli dan Dokumen Revisi.
    Tujuannya adalah untuk memastikan bahwa konten di setiap sub-bab pada Dokumen Revisi tetap sejalan dengan makna dan konteks sub-bab yang sesuai di Dokumen Asli.

    Instruksi:
    1.  Baca dan pahami struktur kedua dokumen. Identifikasi semua bab dan sub-bab (misalnya: "Bab 1", "1.1 Latar Belakang", "2.3 Prosedur Audit").
    2.  **ABAIKAN** bagian-bagian berikut dalam analisis Anda:
        *   Tabel dan isinya.
        *   Gambar dan keterangan gambar.
        *   Daftar pustaka.
        *   Lampiran.
        *   Bagian pendahuluan, definisi, atau daftar istilah jika ada di section/subsection tersendiri.
    3.  Fokus analisis Anda **HANYA** pada isi paragraf dan poin-poin di dalam setiap sub-bab.
    4.  Untuk setiap sub-bab di Dokumen Revisi, bandingkan kontennya dengan sub-bab yang sesuai di Dokumen Asli.
    5.  Identifikasi kalimat atau poin di Dokumen Revisi yang mengalami **perubahan makna yang signifikan**, keluar dari konteks, atau tidak lagi relevan dengan sub-bab aslinya.
    6.  Untuk setiap kalimat/poin yang menyimpang, jelaskan alasannya dalam **dua poin yang jelas**:
        a. **Poin 1: Makna Asli.** Jelaskan makna atau konteks utama dari kalimat ini sebagaimana adanya di Dokumen Asli.
        b. **Poin 2: Ketidaksesuaian & Rekomendasi.** Jelaskan mengapa makna ini tidak lagi relevan atau menyimpang dari konteks di Dokumen Revisi. Berikan saran spesifik tentang apa yang perlu ditambahkan atau diubah agar kalimat ini kembali sejalan dengan konteksnya di Dokumen Revisi.

    Berikut adalah kedua dokumen tersebut:
    ---
    DOKUMEN ASLI:
    ---
    {original_text}
    ---
    DOKUMEN REVISI:
    ---
    {revised_text}
    ---

    Berikan hasil analisis Anda secara eksklusif dalam format JSON array. Setiap objek dalam array harus memiliki tiga kunci:
    1.  "sub_bab_asal": Nama sub-bab (dari Dokumen Revisi) di mana kalimat menyimpang ini ditemukan.
    2.  "kalimat_menyimpang": Kalimat atau poin dari Dokumen Revisi yang maknanya menyimpang.
    3.  "alasan": Penjelasan dalam dua poin. Gunakan format:
        "1. [Penjelasan Makna Asli]\n2. [Penjelasan Ketidaksesuaian & Rekomendasi]"

    Contoh Output:
    [
      {{
        "sub_bab_asal": "1.2 Tujuan Audit",
        "kalimat_menyimpang": "Manajemen wajib memastikan kepatuhan terhadap peraturan perundang-undangan yang berlaku di negara lain.",
        "alasan": "1. Makna asli kalimat ini di dokumen asli adalah pembatasan cakupan hanya pada peraturan perundang-undangan nasional.\n2. Kalimat ini di revisi menjadi tidak relevan karena menambahkan cakupan yurisdiksi internasional yang tidak ada di konteks asli. Untuk menyelaraskan, sebaiknya hapus frasa 'di negara lain' atau ganti dengan 'di Indonesia'."
      }}
    ]

    PENTING: HANYA KELUARKAN JSON ARRAY MURNI. TANPA TEKS PENDAHULU ATAU PENUTUP.
    """
    try:
        response = model.generate_content(prompt)
        # Membersihkan respons dan mencoba parsing JSON
        response_text = response.text.strip()
        response_text = re.sub(r'```json\s*|\s*```', '', response_text)
        
        analysis_result = json.loads(response_text)
        
        # Validasi struktur JSON
        if not isinstance(analysis_result, list):
            raise ValueError("Respons dari AI bukanlah sebuah array JSON.")
            
        return analysis_result

    except json.JSONDecodeError as e:
        print(f"[ERROR] Gagal parsing JSON dari AI: {e}")
        print(f"[ERROR] AI Response was: {response.text if 'response' in locals() else 'N/A'}")
        return [{"sub_bab_asal": "Error", "kalimat_menyimpang": "Gagal memproses respons AI.", "alasan": str(e)}]
    except Exception as e:
        print(f"[ERROR] Terjadi kesalahan saat menganalisis dokumen: {e}")
        return [{"sub_bab_asal": "Error", "kalimat_menyimpang": "Terjadi kesalahan server.", "alasan": str(e)}]


def analyze_context_difference(original_sentence, revised_sentence):
    """
    Menganalisis mengapa konteks kalimat revisi berbeda dari kalimat asli.
    Versi ini lebih sederhana karena nomor halaman disediakan oleh backend.
    """
    if not original_sentence or not revised_sentence:
        return {"alasan": "Tidak cukup data untuk dianalisis.", "kalimat_menyimpang": revised_sentence}

    prompt = f"""
    Anda adalah seorang auditor ahli. Anda diberikan dua dokumen, satu dokumen asli dan satu dokumen lainnya
    Tugas Anda adalah menganalisis mengapa konteks, makna, fokus kalimat / paragraf pada dokumen revisi berbeda dari dokumen asli. Beberapa ketentuannya sebagai berikut
    1. Struktur dari dokumen asli dengan dokumen revisi berbeda, tetapi membahas mengenai hal yang sama.
    2. Anda pelajari terlebih dahulu di dokumen Asli itu maknanya seperti apa secara detail terutama setiap poin poinnya, begitupun juga untuk di dokumen revisi
    3. Tolong Anda periksa secara detail apakah ada perubahan makna yang signifikan pada kalimat di dokumen revisi dibandingkan dokumen aslinya. 
    4. Identifikasi secara teliti semua kalimat yang ada di dokumen revisi yang memiliki perubahan makna signifikan dibandingkan dokumen aslinya.
    5. Berikan penjelasan singkat dan jelas mengapa kalimat tersebut diubah (misalnya: "Menambahkan detail spesifik", "Mengubah fokus dari A ke B", "Memperjelas ambigu", "Mengoreksi fakta").
    6. Jelaskan konteks atau makna utama kalimat asli dan kalimat revisi secara singkat.
    7. Tampilkan alasan yang sangat detail karena ini adalah dokumen Audit sehingga harus sangat teliti dan membuat para Auditor paham mengapa ada perubahan makna pada kalimat tersebut.
    8. Yang ditempilkan pada tabel hasil itu nanti jangan seluruh paragrafnya, tapi cukup kalimat yang mengalami perubahan makna signifikan saja. 

    Kalimat Asli: "{original_sentence}"
    Kalimat Revisi: "{revised_sentence}"

    Analisis dan identifikasi apakah terjadi perubahan makna yang signifikan. Jelaskan alasan perubahannya secara ringkas namun jelas.

    Berikan analisis Anda secara eksklusif dalam format JSON objek dengan dua kunci berikut:
    1. "alasan": Jelaskan secara ringkas mengapa kalimat tersebut diubah. Jika tidak ada perubahan makna signifikan, isi dengan "Tidak ada perubahan makna signifikan."
    2. "kalimat_menyimpang": Salin kembali kalimat revisi yang mengalami perubahan makna.

    Contoh Output:
    {{
        "alasan": "Menambahkan kewajiban hukum yang spesifik.",
        "kalimat_menyimpang": "Manajemen harus memastikan kepatuhan terhadap peraturan perundang-undangan yang berlaku."
    }}

    PENTING: HANYA KELUARKAN OBJEK JSON MURNI. TANPA TEKS PENDAHULU ATAU PENUTUP.
    """

    try:
        response = model.generate_content(prompt)
        response_text = response.text.strip()
        
        # Logika parsing yang sama seperti sebelumnya (lebih tangguh)
        try:
            analysis_result = json.loads(response_text)
        except json.JSONDecodeError:
            print(f"[DEBUG] Gagal parsing JSON langsung. Mencoba ekstraksi. Response: {response_text}")
            match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if match:
                json_string = match.group(0)
                analysis_result = json.loads(json_string)
            else:
                raise ValueError("Tidak dapat menemukan objek JSON yang valid dalam respons AI.")

        # Validasi struktur JSON yang baru
        required_keys = ["alasan", "kalimat_menyimpang"]
        if not all(key in analysis_result for key in required_keys):
            print(f"[DEBUG] Struktur JSON tidak lengkap. Diterima: {analysis_result}")
            raise ValueError("Struktur JSON dari AI tidak lengkap atau tidak sesuai format.")
            
        return analysis_result

    except Exception as e:
        print(f"[ERROR] Terjadi kesalahan saat menganalisis konteks: {e}")
        print(f"[ERROR] Original: {original_sentence}")
        print(f"[ERROR] Revised: {revised_sentence}")
        print(f"[ERROR] AI Response was: {response.text if 'response' in locals() else 'N/A'}")
        
        return {
            "alasan": f"Error: AI gagal memberikan respons yang valid. ({str(e)})",
            "kalimat_menyimpang": revised_sentence
        }

def analyze_document_coherence(full_text):
    if not full_text or full_text.isspace():
        return []

    prompt = f"""
    Anda adalah seorang auditor ahli yang bertugas menganalisis struktur dan koherensi sebuah tulisan.
    Tugas Anda adalah membaca keseluruhan teks berikut dan mengidentifikasi setiap kalimat atau paragraf yang tidak koheren atau keluar dari topik utama di dalam sebuah sub-bagian.
    
    Untuk setiap ketidaksesuaian yang Anda temukan, lakukan hal berikut:
    1. Bacalah mengenai judul dari section atau subsection yang ada pada file tersebut, terutama bacalah isi paragrafnya dan makna setiap kalimatnya
    2. Tentukan topik utama dari setiap section / subsection terutama isi paragrafnya.
    3. Identifikasi kalimat asli yang menyimpang dari topik tersebut.
    4. Berikan saran dengan menghighlight kalimat tersebut untuk ditulis ulang (rewording) agar relevan dan menyatu kembali dengan topik utamanya, ikuti standar KBBI, PUEBI, dan SPOK Bahasa Indonesia.
    5. Jika Anda memiliki asumsi atau catatan tambahan tentang revisi tersebut (seperti "asumsi logis" atau "catatan: ..."), PISAHKAN catatan itu.

    Berikan hasil dalam format yang SANGAT KETAT.
    Format 1 (Tanpa Catatan):
    [TOPIK UTAMA] topik utama -> [TEKS ASLI] kalimat asli -> [SARAN REVISI] versi kalimat yang sudah diperbaiki

    Format 2 (Dengan Catatan):
    [TOPIK UTAMA] topik utama -> [TEKS ASLI] kalimat asli -> [SARAN REVISI] versi kalimat yang sudah diperbaiki -> [CATATAN] asumsi atau catatan Anda

    Contoh Format 2:
    [TOPIK UTAMA] Rencana Kerja Tahunan -> [TEKS ASLI] Penyebab utamanya adalah... -> [SARAN REVISI] Penyebab utamanya adalah... -> [CATATAN] Asumsi logis dari konteks.

    Jika seluruh dokumen sudah koheren dan tidak ada masalah, kembalikan teks: "TIDAK ADA MASALAH KOHERENSI"

    Teks:
    ---
    {full_text}
    """
    try:
        response = model.generate_content(prompt)
        pattern = re.compile(
            r"\[TOPIK UTAMA\]\s*(.*?)\s*->\s*\[TEKS ASLI\]\s*(.*?)\s*->\s*\[SARAN REVISI\]\s*(.*?)\s*(?:->\s*\[CATATAN\]\s*(.*?)\s*)?(\n|$)", 
            re.IGNORECASE | re.DOTALL
        )
        found_issues = pattern.findall(response.text)
        
        results = []
        for topik, asli, saran, catatan, _ in found_issues:
            results.append({
                "topik": topik.strip(), 
                "asli": asli.strip(), 
                "saran": saran.strip(), 
                "catatan": catatan.strip() if catatan else ""
            })
        return results
        
    except Exception as e:
        print(f"Terjadi kesalahan saat menghubungi AI: {e}")
        return [{"topik": "ERROR", "asli": str(e), "saran": "Gagal menghubungi API", "catatan": ""}]

def get_structural_recommendations(full_text):
    if not full_text or full_text.isspace():
        return []
    # PROMPT RESTRUKTURISASI YANG TELAH DIKOREKSI DARI KARAKTER AMBIGU
    prompt = f"""
    Anda adalah seorang auditor ahli yang bertugas untuk melakukan analisis terhadap dokumen. Tugas Anda adalah menemukan paragraf yang terkesan 'salah tempat' dan memberikan saran di bagian mana seharusnya paragraf tersebut berada saat ini (lokasi asli).

    Untuk setiap paragraf yang terdeteksi, Anda harus:
    1. Bacalah semua dokumennya terlebih dahulu, temukan ide-ide utama di setiap paragraf, dan merevisi jika perlu.
    2. Pada saat Anda membaca dokumennya, tolong identifikasi teks lengkap dari paragraf yang berada tidak pada tempatnya.
    3. Tentukan di sub-bab atau sub-bab mana paragraf itu berada saat ini (lokasi asli).
    4. Berikan rekomendasi di sub-bab atau sub-bab mana paragraf tersebut seharusnya diletakkan agar lebih koheren dan logis.
    5. Kalau ada bagian yang harus dipindahkan ke Ringkasan Eksekutif, itu tidak perlu dimasukkan ke dalam tabel.
    6. Kalau ada kata yang merupakan bahasa inggris, biarkan saja dan tidak perlu diitalic.
    7. Kalau ada kata yang tidak baku sesuai dengan standar KBBI, harap Anda perbaiki saja. Sehingga kata tersebut menjadi kata baku.
    8. Pada bagian lampiran, tidak perlu di cek/dicek untuk dipindahkan ke bagian lainnya karena itu sudah fixed / sudah benar

    Berikan hasil dalam format JSON yang berisi sebuah list. Setiap objek harus memiliki tiga kunci: "misplaced_paragraph", "original_section", dan "recommended_section".

    Contoh Format JSON:
    [
      {{
        "misplaced_paragraph": "Selain itu, audit internal juga bertugas memeriksa laporan keuangan setiap kuartal...",
        "original_section": "Bab 2.1: Prosedur Whistleblowing",
        "recommended_section": "Bab 4.2: Peran Audit Internal"
      }}
    ]

    Berikut adalah teks yang harus Anda periksa:
    ---
    {full_text}
    """
    try:
        response = model.generate_content(prompt)
        # FIX: Mengganti karakter 'long dash' (U+2014) jika ada
        cleaned_response = re.sub(r'[—–]', '-', response.text.strip()) # Tambahkan penanganan untuk U+2014 dan U+2013
        cleaned_response = re.sub(r'```json\s*|\s*```', '', cleaned_response)
        return json.loads(cleaned_response)
    except Exception as e:
        print(f"Failed to Generate Response from AI: {e}")
        # Fallback jika AI tidak mengembalikan JSON
        return [{"misplaced_paragraph": "Error: " + str(e), "original_section": "Gagal menghubungi API", "recommended_section": "Periksa prompt Anda."}]

def generate_revised_docx(file_bytes, errors):
    doc = docx.Document(io.BytesIO(file_bytes))
    
    for error in reversed(errors):
        # Sesuaikan key ini berdasarkan apa yang dikirim
        salah = error.get("salah") or error.get("Kata/Frasa Salah")
        benar = error.get("benar") or error.get("Perbaikan Sesuai KBBI")
        
        if not salah or not benar:
            continue
            
        for para in doc.paragraphs:
            if salah in para.text:
                para.text = para.text.replace(salah, benar) 
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()

def generate_highlighted_docx(file_bytes, errors):
    doc = docx.Document(io.BytesIO(file_bytes))
    
    # Sesuaikan key ini
    unique_salah = set(e.get("salah") or e.get("Kata/Frasa Salah") for e in errors if e.get("salah") or e.get("Kata/Frasa Salah"))
    
    for para in doc.paragraphs:
        for term in unique_salah:
            if term and term.lower() in para.text.lower():
                full_text = para.text
                para.clear()
                parts = re.split(f'({re.escape(term)})', full_text, flags=re.IGNORECASE)
                for part in parts:
                    if part:
                        run = para.add_run(part)
                        if part.lower() == term.lower():
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()

def create_zip_archive(revised_data, highlighted_data, original_filename):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        zip_file.writestr(f"revisi_{original_filename}", revised_data)
        zip_file.writestr(f"highlight_{original_filename}", highlighted_data)
    return zip_buffer.getvalue()

def parse_flexible_date(date_str):
    """Mencoba parsing tanggal dari berbagai format."""
    if not date_str:
        return None
    try:
        # Coba parsing dengan waktu (format datetime-local)
        return datetime.datetime.strptime(date_str, '%Y-%m-%dT%H:%M')
    except ValueError:
        try:
            # Jika gagal, coba parsing tanggal saja (format date)
            return datetime.datetime.strptime(date_str, '%Y-%m-%d')
        except ValueError:
            raise ValueError("Format tanggal tidak valid. Gunakan format YYYY-MM-DD atau YYYY-MM-DDTHH:MM.")

def extract_paragraphs(file_bytes):
    try:
        source_stream = io.BytesIO(file_bytes)
        doc = docx.Document(source_stream)
        return [p.text for p in doc.paragraphs if p.text.strip() != ""]
    except Exception as e:
        raise ValueError(f"Gagal membaca file docx: {e}")

def extract_paragraphs_from_text(full_text):
    if not full_text:
        return []
    # Memisahkan berdasarkan dua baris baru, lalu membersihkan whitespace
    paragraphs = [p.strip() for p in full_text.split('\n\n') if p.strip()]
    return paragraphs

def find_word_diff(original_para, revised_para):
    matcher = difflib.SequenceMatcher(None, original_para.split(), revised_para.split())
    diffs = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'replace' or tag == 'insert':
            diffs.append(" ".join(revised_para.split()[j1:j2]))
    return ", ".join(diffs) if diffs else "Perubahan Minor"

def create_comparison_docx(df):
    doc = Document() # Pastikan 'from docx import Document' ada di atas
    doc.add_heading('Hasil Perbandingan Dokumen', level=1)
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = col_name
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col_name in enumerate(df.columns):
            row_cells[i].text = str(row[col_name])
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()

def calculate_task_status(start_time, deadline, end_time):
    if end_time:
        return 'done'
    
    if deadline and start_time and start_time.date() > deadline:
        return 'overdue'
    
    # Jika tidak ada kondisi yang terpenuhi, statusnya on_progress
    # Ini akan menangani tugas manual yang belum selesai dan belum overdue
    return 'on_progress'

def create_recommendation_highlight_docx(file_bytes, recommendations):
    doc = docx.Document(io.BytesIO(file_bytes))
    
    # Sesuaikan key ini
    misplaced_paragraphs = [rec.get("misplaced_paragraph") or rec.get("Paragraf yang Perlu Dipindah") for rec in recommendations]
    
    for para in doc.paragraphs:
        if para.text.strip() in [p.strip() for p in misplaced_paragraphs if p]:
            for run in para.runs:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()


def _get_word_diff_structure(original_para, revised_para):
    original_words = original_para.split()
    revised_words = revised_para.split()
    matcher = difflib.SequenceMatcher(None, original_words, revised_words)
    
    structured_output = []
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            structured_output.append({
                "text": " ".join(revised_words[j1:j2]),
                "changed": False
            })
        elif tag == 'replace' or tag == 'insert':
            structured_output.append({
                "text": " ".join(revised_words[j1:j2]),
                "changed": True
            })
    
    # Tambahkan spasi di akhir setiap bagian agar tidak menempel
    for item in structured_output:
        item['text'] += ' '
        
    return structured_output

def _analyze_comparison(file1, file2):
    file1_bytes = file1.read()
    file2_bytes = file2.read()
    file1.seek(0)
    file2.seek(0)

    file1_ext = file1.filename.split('.')[-1].lower()
    file2_ext = file2.filename.split('.')[-1].lower()

    pages1 = _extract_text_with_pages(file1_bytes, file1_ext)
    pages2 = _extract_text_with_pages(file2_bytes, file2_ext)
    
    original_data = extract_sentences_with_pages(pages1)
    revised_data = extract_sentences_with_pages(pages2)
    
    original_sentences = [item['sentence'] for item in original_data]
    revised_sentences = [item['sentence'] for item in revised_data]
    
    comparison_results = []
    matcher = difflib.SequenceMatcher(None, original_sentences, revised_sentences)
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'replace':
            len_original = i2 - i1
            len_revised = j2 - j1
            
            for i in range(min(len_original, len_revised)):
                original_sentence = original_sentences[i1 + i]
                revised_sentence = revised_sentences[j1 + i]
                
                # Ambil nomor halaman dari data yang sudah diproses
                revised_page = revised_data[j1 + i]['page']
                
                word_diff_text = find_word_diff(original_sentence, revised_sentence)
                revised_structured = _get_word_diff_structure(original_sentence, revised_sentence)
                
                comparison_results.append({
                    "Kalimat Awal": original_sentence,
                    "Kalimat Revisi": revised_structured,
                    "Kata yang Direvisi": word_diff_text,
                    "Halaman": f"Halaman {revised_page}"
                })
    return comparison_results

@app.route('/', methods=['GET', 'POST']) 
def login():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard')) 
    
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = User.query.filter_by(username=username).first()
        
        if user and user.check_password(password):
            login_user(user)
            return redirect(url_for('dashboard'))
        else:
            flash('Username atau password salah. Harap periksa kembali kesesuaian username dan password Anda.') 
            
    return render_template('login.html')

@app.route('/logout')
@login_required 
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/dashboard') 
@login_required 
def dashboard(): 
    """
    REVISI: Mengambil data folder (objek) dari get_user_folders.
    """
    folders = []
    try:
        folders = get_user_folders() # Ini sekarang mengembalikan list of dicts
    except Exception as e:
        print(f"Error memuat folder untuk dashboard: {e}")
        flash("Gagal memuat struktur folder Anda.")
    
    # Mengirim data 'folders' (list of dicts) ke template
    return render_template('index.html', 
                            username=current_user.username, 
                            label=current_user.label,
                            folders=folders,
                            current_user_id=current_user.id) # TAMBAHAN: Kirim user ID

# --- API BARU UNTUK FOLDER ---
@app.route('/api/create_folder', methods=['POST'])
@login_required 
def api_create_folder():
    """Membuat folder baru di sistem file."""
    data = request.json
    name = data.get('name')

    if not name:
        return jsonify({"error": "Nama folder tidak boleh kosong."}), 400

    try:
        clean_name = create_user_folder(name)
        return jsonify({
            "status": "success",
            "folder_name": clean_name
        }), 201
    except ValueError as e:
        # Kesalahan yang disengaja (misal: nama duplikat)
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        # Kesalahan server (misal: izin folder)
        print(f"Kesalahan Server saat membuat folder: {e}")
        return jsonify({"error": "Gagal membuat folder di server."}), 500

@app.route('/api/list_folders', methods=['GET'])
@login_required 
def api_list_folders():
    """Mendaftarkan semua folder yang dimiliki pengguna."""
    print(f"Mencoba mengambil folder untuk user: {current_user.id}") # DEBUG
    try:
        folders = get_user_folders() # Ini sekarang mengembalikan list of dicts
        print(f"Ditemukan folder: {folders}") # DEBUG
        return jsonify(folders)
    except Exception as e:
        print(f"Kesalahan Server saat mengambil list folder: {e}")
        return jsonify({"error": "Gagal memuat folder: " + str(e)}), 500

@app.route('/api/delete_folder', methods=['POST'])
@login_required
def api_delete_folder():
    """Menghapus folder dan seluruh isinya."""
    data = request.json
    folder_name = data.get('folder_name')

    if not folder_name or '..' in folder_name or '~' in folder_name:
        return jsonify({"error": "Nama folder tidak valid."}), 400
    
    try:
        user_root = get_user_root_folder()
        if not user_root:
            return jsonify({"error": "User tidak ditemukan."}), 401
            
        folder_path = os.path.join(user_root, folder_name)

        # Keamanan: Pastikan path aman dan berada di dalam user_root
        if not os.path.isdir(folder_path) or not os.path.abspath(folder_path).startswith(os.path.abspath(user_root)):
             return jsonify({"error": "Folder tidak ditemukan atau akses ditolak."}), 404

        # Hapus folder dan semua isinya secara rekursif
        shutil.rmtree(folder_path)
        
        # TAMBAHAN: Hapus juga catatan 'share' jika ada
        SharedFolder.query.filter_by(owner_id=current_user.id, folder_name=folder_name).delete()
        db.session.commit()
        
        return jsonify({"status": "success", "message": f"Folder '{folder_name}' berhasil dihapus."}), 200

    except Exception as e:
        db.session.rollback()
        print(f"Error saat menghapus folder: {e}")
        return jsonify({"error": f"Gagal menghapus folder: {e}"}), 500


@app.route('/api/save_results', methods=['POST'])
@login_required 
def api_save_results():
    """Menyimpan hasil analisis ke file JSON di dalam folder."""
    data = request.json
    folder_name = data.get('folder_name')
    feature_type = data.get('feature_type')
    results_data = data.get('results_data')
    original_filename = data.get('original_filename', 'untitled_analysis')
    owner_id = data.get('owner_id', current_user.id)
    # >>>>>> AMBIL DATA AKSI DARI REQUEST <<<<<<
    actions_data = data.get('actions_data', {})

    if not folder_name or not feature_type or not results_data:
        return jsonify({"error": "Data folder, fitur, atau hasil kosong."}), 400

    try:
        # --- REVISI: Tentukan root folder berdasarkan owner_id ---
        target_user_id_str = str(owner_id)
        if not target_user_id_str.isalnum():
             return jsonify({"error": "Owner ID tidak valid."}), 400
        
        user_root = os.path.join(app.config['UPLOAD_FOLDER'], target_user_id_str)
        
        # Keamanan: Pastikan folder_name tidak mengandung '..'
        if '..' in folder_name:
             return jsonify({"error": "Nama folder tidak valid."}), 400
             
        folder_path = os.path.join(user_root, folder_name)

        # --- REVISI: Cek Izin ---
        is_owner = (str(current_user.id) == target_user_id_str)
        is_shared_to_me = SharedFolder.query.filter_by(
            owner_id=owner_id, 
            folder_name=folder_name, 
            shared_with_id=current_user.id
        ).first()

        if not is_owner and not is_shared_to_me:
            return jsonify({"error": "Akses ditolak untuk menyimpan ke folder ini."}), 403
        
        if not os.path.isdir(folder_path):
            return jsonify({"error": "Folder tidak ditemukan."}), 404
        # --- AKHIR REVISI Izin ---

        # 2. Tentukan nama file
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        # Bersihkan nama file asli
        clean_orig_name = re.sub(r'[^\w\s-]', '', original_filename.split('.')[0]).strip()
        clean_orig_name = re.sub(r'[-\s]+', '_', clean_orig_name)
        
        save_filename = f"{timestamp}_{feature_type}_{clean_orig_name}.json"
        save_path = os.path.join(folder_path, save_filename)

        # 3. Simpan data (JSON)
        with open(save_path, 'w', encoding='utf-8') as f:
            json.dump(results_data, f, ensure_ascii=False, indent=2)

        # >>>>>> SIMPAN DATA AKSI KE DATABASE <<<<<<
        # Loop melalui data aksi yang diterima dari frontend
        for row_id_str, action_data in actions_data.items():
            try:
                row_id = int(row_id_str) # Konversi ID baris dari string ke integer
                
                # Cek apakah aksi untuk baris ini sudah ada
                action = RowAction.query.filter_by(
                    owner_id=owner_id,
                    folder_name=folder_name,
                    filename=save_filename, # Gunakan nama file yang baru disimpan
                    row_id=row_id
                ).first()

                if action:
                    # Jika sudah ada, update data-nya
                    action.is_ganti = action_data.get('is_ganti', False)
                    action.pic_user_id = action_data.get('pic_user_id')
                else:
                    # Jika belum ada, buat entri baru
                    action = RowAction(
                        owner_id=owner_id,
                        folder_name=folder_name,
                        filename=save_filename,
                        row_id=row_id,
                        is_ganti=action_data.get('is_ganti', False),
                        pic_user_id=action_data.get('pic_user_id')
                    )
                    db.session.add(action)
            except Exception as e:
                # Jika terjadi error pada salah satu baris, batalkan semua transaksi
                print(f"Gagal memproses aksi untuk baris {row_id_str}: {e}")
                raise e # Lempar error ke blok try-catch utama
        
        # Commit semua perubahan ke database sekali saja setelah loop selesai
        db.session.commit()
        # >>>>>> AKHIR SIMPAN DATA AKSI <<<<<<

        return jsonify({
            "status": "success",
            "message": f"Hasil analisis dan status aksi tersimpan di folder {folder_name}."
        }), 201
    except Exception as e:
        # Jika ada error, rollback perubahan database
        db.session.rollback()
        print(f"Kesalahan Server saat menyimpan hasil: {e}")
        return jsonify({"error": "Gagal menyimpan hasil: " + str(e)}), 500
    
# --- TAMBAHAN BARU: API UNTUK FITUR SHARE & RIWAYAT ---

@app.route('/api/get_all_users', methods=['GET'])
@login_required
def api_get_all_users():
    """Mengambil daftar semua user KECUALI diri sendiri."""
    try:
        users = User.query.all()
        # MODIFIKASI: Kirim juga 'label' (tag)
        user_list = [{"id": user.id, "username": user.username, "label": user.label} for user in users]
        return jsonify(user_list)
    except Exception as e:
        print(f"Error get all users: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/share_folder', methods=['POST'])
@login_required
def api_share_folder():
    """Berbagi folder dengan user lain."""
    data = request.json
    folder_name = data.get('folder_name')
    share_with_user_ids = data.get('share_with_user_ids')

    if not folder_name or not share_with_user_ids:
        return jsonify({"error": "Data tidak lengkap (folder atau user)."}), 400
    
    if not isinstance(share_with_user_ids, list):
         return jsonify({"error": "Format user ID harus berupa list."}), 400

    success_names = []  # <<< PERUBAHAN: Ganti count dengan list nama
    skipped_count = 0
    errors = []

    for user_id_str in share_with_user_ids: 
        try:
            user_id = int(user_id_str) # KONVERSI WAJIB
        except ValueError:
            errors.append(f"User ID tidak valid: {user_id_str}")
            continue
    
        user_to_share = User.query.get(user_id)
        if not user_to_share:
            errors.append(f"Pengguna dengan ID {user_id} tidak ditemukan.")
            continue

        # Cek apakah sudah pernah di-share
        existing_share = SharedFolder.query.filter_by(
            owner_id=current_user.id,
            folder_name=folder_name,
            shared_with_id=user_id
        ).first()

        if existing_share:
            skipped_count += 1
            continue

        try:
            new_share = SharedFolder(
                owner_id=current_user.id,
                folder_name=folder_name,
                shared_with_id=user_id,
                folder_type='general' 
            )
            db.session.add(new_share)
            success_names.append(user_to_share.username) # <<< PERUBAHAN: Tambahkan nama
        
        except Exception as e:
            db.session.rollback() 
            errors.append(f"Gagal share ke {user_to_share.username}: {e}")
            continue 

    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": f"Gagal menyimpan ke database: {e}"}), 500

    return jsonify({
        "status": "success",
        "message": f"Berhasil di-share ke {len(success_names)} user.",
        "success_names": success_names,  # <<< PERUBAHAN: Mengembalikan list nama
        "skipped_count": skipped_count,
        "errors": errors
    }), 201
    

@app.route('/api/folder_history/<int:owner_id>/<folder_name>', methods=['GET'])
@login_required
def api_folder_history(owner_id, folder_name):
    """
    PERBAIKAN UNTUK MASALAH 1: Mengambil riwayat file dari folder.
    Kita butuh owner_id untuk folder yang di-share.
    """
    if not folder_name or '..' in folder_name:
        return jsonify({"error": "Nama folder tidak valid."}), 400
    
    try:
        # Tentukan path folder. Jika owner_id adalah ID kita, gunakan root kita.
        # Jika beda, itu adalah folder yang di-share.
        
        target_user_id_str = str(owner_id)
        
        # Validasi keamanan dasar
        if not target_user_id_str.isalnum():
             return jsonify({"error": "Owner ID tidak valid."}), 400
        
        folder_root = os.path.join(app.config['UPLOAD_FOLDER'], target_user_id_str)
        folder_path = os.path.join(folder_root, folder_name)

        # Cek izin: Apakah kita pemilik ATAU folder ini di-share ke kita?
        is_owner = (str(current_user.id) == target_user_id_str)
        is_shared_to_me = SharedFolder.query.filter_by(
            owner_id=owner_id, 
            folder_name=folder_name, 
            shared_with_id=current_user.id
        ).first()

        if not is_owner and not is_shared_to_me:
            return jsonify({"error": "Akses ditolak."}), 403
            
        if not os.path.isdir(folder_path):
            return jsonify({"error": "Folder tidak ditemukan."}), 404

        results = []
        for filename in os.listdir(folder_path):
            if filename.endswith('.json'):
                try:
                    # Parse nama file: 20251106_223800_proofreading_NamaFileAsli.json
                    parts = filename.replace('.json', '').split('_', 3)
                    timestamp_str = f"{parts[0]}_{parts[1]}"
                    timestamp = datetime.datetime.strptime(timestamp_str, "%Y%m%d_%H%M%S")
                    feature_type = parts[2]
                    original_name = parts[3] if len(parts) > 3 else "N/A"
                    
                    results.append({
                        "filename": filename,
                        "feature_type": feature_type,
                        "timestamp": timestamp.strftime("%d %b %Y, %H:%M"),
                        "original_name": original_name.replace("_", " ")
                    })
                except Exception as e:
                    print(f"Gagal parse nama file {filename}: {e}")
                    # Jika parse gagal, tampilkan mentah
                    results.append({
                        "filename": filename,
                        "feature_type": "N/A",
                        "timestamp": "N/A",
                        "original_name": "N/A"
                    })
        
        # Urutkan berdasarkan timestamp terbaru
        results.sort(key=lambda x: x.get("timestamp"), reverse=True)
        
        return jsonify(results)

    except Exception as e:
        print(f"Error saat mengambil riwayat folder: {e}")
        return jsonify({"error": str(e)}), 500

# --- TAMBAHAN BARU: API UNTUK MELIHAT ISI FILE JSON ---
@app.route('/api/get_result_file', methods=['POST'])
@login_required
def api_get_result_file():
    """Membaca dan mengembalikan isi file JSON yang disimpan, beserta status aksinya."""
    data = request.json
    folder_name = data.get('folder_name')
    filename = data.get('filename')
    owner_id = data.get('owner_id')

    if not folder_name or not filename or not owner_id:
        return jsonify({"error": "Data tidak lengkap."}), 400
    
    if '..' in folder_name or '..' in filename:
        return jsonify({"error": "Nama file/folder tidak valid."}), 400

    try:
        # Tentukan path
        target_user_id_str = str(owner_id)
        if not target_user_id_str.isalnum():
             return jsonify({"error": "Owner ID tidak valid."}), 400
        
        folder_root = os.path.join(app.config['UPLOAD_FOLDER'], target_user_id_str)
        file_path = os.path.join(folder_root, folder_name, filename)

        # Cek izin: Apakah kita pemilik ATAU folder ini di-share ke kita?
        is_owner = (str(current_user.id) == target_user_id_str)
        is_shared_to_me = SharedFolder.query.filter_by(
            owner_id=owner_id, 
            folder_name=folder_name, 
            shared_with_id=current_user.id
        ).first()

        if not is_owner and not is_shared_to_me:
            return jsonify({"error": "Akses ditolak."}), 400
            
        if not os.path.isfile(file_path):
            return jsonify({"error": "File tidak ditemukan."}), 404

        # 1. Baca file JSON hasil analisis
        with open(file_path, 'r', encoding='utf-8') as f:
            json_data = json.load(f)
            
        # 2. >>>>>> TAMBAHAN: Ambil data aksi per baris dari database <<<<<<
        row_actions = RowAction.query.filter_by(
            owner_id=owner_id,
            folder_name=folder_name,
            filename=filename
        ).all()

        actions_data = {
            action.row_id: {
                'is_ganti': action.is_ganti, 
                'pic_user_id': action.pic_user_id
            } for action in row_actions
        }
        return jsonify({
            "status": "success",
            "data": json_data,
            "actions": actions_data
        })
    
        response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate'
        response.headers['Pragma'] = 'no-cache'
        response.headers['Expires'] = '0'
        
        return response

    except Exception as e:
        print(f"Error saat membaca file result: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/delete_result', methods=['POST'])
@login_required
def api_delete_result():
    
    """Menghapus satu file hasil analisis (JSON) dari folder."""
    data = request.json
    folder_name = data.get('folder_name')
    filename = data.get('filename')
    owner_id = data.get('owner_id') # Kita butuh ini untuk keamanan

    if not folder_name or not filename or not owner_id:
        return jsonify({"error": "Data tidak lengkap."}), 400
    
    if '..' in folder_name or '..' in filename:
        return jsonify({"error": "Nama file/folder tidak valid."}), 400

    try:
        # Cek izin: Hanya pemilik folder yang boleh menghapus file
        if str(current_user.id) != str(owner_id):
             return jsonify({"error": "Akses ditolak. Hanya pemilik folder yang bisa menghapus file."}), 403

        user_root = get_user_root_folder() # Root folder milik kita
        file_path = os.path.join(user_root, folder_name, filename)

        # Keamanan: Pastikan path aman
        if not os.path.isfile(file_path) or not os.path.abspath(file_path).startswith(os.path.abspath(user_root)):
            return jsonify({"error": "File tidak ditemukan atau akses ditolak."}), 404
        
        os.remove(file_path)
        
        return jsonify({"status": "success", "message": f"File '{filename}' berhasil dihapus."}), 200

    except Exception as e:
        print(f"Error saat menghapus file: {e}")
        return jsonify({"error": f"Gagal menghapus file: {e}"}), 500

# --- Endpoint API Fitur (Dipertahankan) ---
@app.route('/api/proofread/analyze', methods=['POST'])
@login_required 
def api_proofread_analyze():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    
    try:
        document_pages = _get_text_from_flask_file(file)
        # Reset file pointer setelah dibaca (jika _get_text_from_flask_file tidak melakukannya)
        file.seek(0)
        all_errors = []
        for page in document_pages:
            found_errors_on_page = proofread_with_gemini(page['teks'])
            for error in found_errors_on_page:
                all_errors.append({
                    "Kata/Frasa Salah": error['salah'],
                    "Perbaikan Sesuai KBBI": error['benar'],
                    "Pada Kalimat": error['kalimat'],
                    "Ditemukan di Halaman": page['halaman']
                })
        
        return jsonify(all_errors)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def _generate_proofread_files(file, file_bytes):
    """Helper internal untuk download, menjalankan analisis lagi."""
    file_extension = file.filename.split('.')[-1].lower()
    document_pages = _extract_text_with_pages(file_bytes, file_extension)
    
    all_errors = []
    for page in document_pages:
        found_errors_on_page = proofread_with_gemini(page['teks'])
        all_errors.extend(found_errors_on_page) 

    revised_data = generate_revised_docx(file_bytes, all_errors)
    highlighted_data = generate_highlighted_docx(file_bytes, all_errors)
    
    return revised_data, highlighted_data, file.filename

@app.route('/api/proofread/download/revised', methods=['POST'])
@login_required 
def api_proofread_download_revised():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    file_bytes = file.read() # Baca sekali
    file.seek(0) # Reset pointer
    
    try:
        revised_data, _, filename = _generate_proofread_files(file, file_bytes)
        
        return send_file(
            io.BytesIO(revised_data),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"revisi_{filename}"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/proofread/download/highlighted', methods=['POST'])
@login_required 
def api_proofread_download_highlighted():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    file_bytes = file.read() # Baca sekali
    file.seek(0)
    
    try:
        _, highlighted_data, filename = _generate_proofread_files(file, file_bytes)
        
        return send_file(
            io.BytesIO(highlighted_data),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"highlight_{filename}"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/proofread/download/zip', methods=['POST'])
@login_required 
def api_proofread_download_zip():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    file_bytes = file.read() # Baca sekali
    file.seek(0)
    
    try:
        revised_data, highlighted_data, filename = _generate_proofread_files(file, file_bytes)
        zip_data = create_zip_archive(revised_data, highlighted_data, filename)
        
        return send_file(
            io.BytesIO(zip_data),
            mimetype='application/zip',
            as_attachment=True,
            download_name=f"hasil_proofread_{filename.split('.')[0]}.zip"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/compare/analyze_advanced', methods=['POST'])
@login_required 
def api_compare_analyze_advanced():
    if 'file1' not in request.files or 'file2' not in request.files:
        return jsonify({"error": "Butuh dua file"}), 400
    
    file1 = request.files['file1']
    file2 = request.files['file2']
    
    try:
        full_text1 = _get_full_text_from_file(file1)
        full_text2 = _get_full_text_from_file(file2)
        
        # Panggil fungsi analisis yang sudah direvisi di atas
        comparison_results_from_ai = analyze_document_by_section(full_text1, full_text2)
        
        # MODIFIKASI 1: Batasi hasil ke 100 teratas
        limited_results = comparison_results_from_ai[:100]

        # MODIFIKASI 2: Pemetaan dan format ulang data
        final_results = []
        for item in limited_results:
            sub_bab_asli = item.get("sub_bab_asal", "N/A")
            
            try:
                # Format nama sub-bab: Ambil teks setelah ':'
                nama_sub_bab = sub_bab_asli.split(':', 1)[1].strip()
            except (IndexError, AttributeError):
                nama_sub_bab = sub_bab_asli
            
            final_results.append({
                "Sub-bab Asal": nama_sub_bab,
                "Kalimat yang Menyimpang di dokumen lainnya": item.get("kalimat_menyimpang", "N/A"),
                "Alasan": item.get("alasan", "N/A")
            })
        
        return jsonify(final_results)
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/compare/analyze', methods=['POST'])
@login_required 
def api_compare_analyze():
    if 'file1' not in request.files or 'file2' not in request.files:
        return jsonify({"error": "Butuh dua file"}), 400
    
    file1 = request.files['file1']
    file2 = request.files['file2']
    
    try:
        # Kita tidak perlu membaca bytes di sini, langsung kirim objek file
        results = _analyze_comparison(file1, file2)
        return jsonify(results)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/compare/download', methods=['POST'])
@login_required 
def api_compare_download():
    if 'file1' not in request.files or 'file2' not in request.files:
        return jsonify({"error": "Butuh dua file"}), 400
    
    file1 = request.files['file1']
    file2 = request.files['file2']
    
    try:
        original_paras = extract_paragraphs(file1.read())
        revised_paras = extract_paragraphs(file2.read())
        comparison_results = []
        matcher = difflib.SequenceMatcher(None, original_paras, revised_paras)
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'replace':
                len_original = i2 - i1
                len_revised = j2 - j1
                for i in range(min(len_original, len_revised)):
                    original_para = original_paras[i1 + i]
                    revised_para = revised_paras[j1 + i]
                    word_diff = find_word_diff(original_para, revised_para)
                    comparison_results.append({
                        "Kalimat Awal": original_para,
                        "Kalimat Revisi": revised_para,
                        "Kata yang Direvisi": word_diff,
                    })
        
        if not comparison_results:
             return jsonify({"error": "Tidak ada perbedaan untuk diunduh"}), 400

        df_comparison = pd.DataFrame(comparison_results)
        docx_data = create_comparison_docx(df_comparison)
        
        return send_file(
            io.BytesIO(docx_data),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"perbandingan_{file1.filename}"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/coherence/analyze', methods=['POST'])
@login_required 
def api_coherence_analyze():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    
    try:
        full_text = _get_full_text_from_file(file)
        issues_from_gemini = analyze_document_coherence(full_text)
        
        processed_issues = []
        for issue in issues_from_gemini:
            asli_text = issue['asli']
            saran_text = issue['saran']
            
            saran_structured = _get_word_diff_structure(asli_text, saran_text)
            
            processed_issues.append({
                "topik": issue['topik'],
                "asli": asli_text,
                "saran": saran_structured,
                "catatan": issue['catatan']
            })
            
        return jsonify(processed_issues)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

def _analyze_restructure(file):
    full_text = _get_full_text_from_file(file)
    recommendations = get_structural_recommendations(full_text)
    processed_results = []
    for rec in recommendations:
        processed_results.append({
            "Paragraf yang Perlu Dipindah": rec.get("misplaced_paragraph"),
            "Lokasi Asli": rec.get("original_section"),
            "Saran Lokasi Baru": rec.get("recommended_section")
        })
    return processed_results

@app.route('/api/restructure/analyze', methods=['POST'])
@login_required 
def api_restructure_analyze():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    
    try:
        results = _analyze_restructure(file)
        return jsonify(results)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/restructure/download', methods=['POST'])
@login_required 
def api_restructure_download():
    if 'file' not in request.files:
        return jsonify({"error": "Tidak ada file"}), 400
    file = request.files['file']
    file_bytes = file.read() # Baca sekali
    file.seek(0) # Reset
    
    try:
        # Gunakan file_bytes untuk ekstraksi teks, bukan 'file'
        full_text = "\n".join([p['teks'] for p in _extract_text_with_pages(file_bytes, file.filename.split('.')[-1].lower())])
        recommendations = get_structural_recommendations(full_text)
        processed_results = []
        for rec in recommendations:
            # Pastikan key di sini sesuai dengan output get_structural_recommendations
            processed_results.append({
                "Paragraf yang Perlu Dipindah": rec.get("misplaced_paragraph") 
            })

        if not processed_results or "Error:" in processed_results[0].get("Paragraf yang Perlu Dipindah", ""):
             return jsonify({"error": "Tidak ada rekomendasi valid untuk diunduh"}), 400

        highlighted_data = create_recommendation_highlight_docx(file_bytes, processed_results)
        
        return send_file(
            io.BytesIO(highlighted_data),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"highlight_rekomendasi_{file.filename}"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
# >>>>>> REVISION START: Route Komentar yang Diperbarui <<<<<<
@app.route('/api/get_comments', methods=['POST'])
@login_required
def api_get_comments():
    data = request.json
    folder_name = data.get('folder_name')
    filename = data.get('filename')

    if not folder_name or not filename:
        return jsonify({"error": "Data file tidak lengkap."}), 400

    # Ambil semua komentar untuk file ini, diurutkan dari yang terlama
    all_comments = Comment.query.filter_by(
        folder_name=folder_name,
        filename=filename
    ).order_by(Comment.timestamp.asc()).all()

    # Buat dictionary untuk memetakan ID ke objek komentar
    comments_by_id = {c.id: c for c in all_comments}
    
    # Pisahkan komentar utama (yang tidak punya parent)
    top_level_comments = [c for c in all_comments if c.parent_id is None]

    # Fungsi rekursif untuk membangun struktur bersarang
    def build_comment_tree(comment):
        comment_dict = {
            'id': comment.id,
            'row_id': comment.row_id,
            'username': comment.username,
            'text': comment.text,
            'timestamp': comment.timestamp.isoformat(),
            'replies': []
        }
        # Cari semua balasan untuk komentar ini
        for reply in all_comments:
            if reply.parent_id == comment.id:
                comment_dict['replies'].append(build_comment_tree(reply))
        return comment_dict

    # Bangun struktur pohon untuk setiap komentar utama
    nested_comments = [build_comment_tree(c) for c in top_level_comments]

    return jsonify(nested_comments), 200

@app.route('/add_comment', methods=['POST'])
@login_required 
def add_comment():
    data = request.get_json()
    folder_name = data.get('folderName')
    filename = data.get('fileName')
    row_id = data.get('rowId')
    text = data.get('text')
    parent_id = data.get('parentId') # Bisa None jika ini komentar utama

    if not folder_name or not filename or not row_id or not text:
        return jsonify({'status': 'error', 'message': 'Data komentar tidak lengkap.'}), 400

    try:
        new_comment = Comment(
            owner_id=current_user.id,
            folder_name=folder_name,
            filename=filename,
            row_id=row_id,
            username=current_user.username,
            text=text,
            # >>>>>> TAMBAHKAN BARIS INI <<<<<<
            parent_id=parent_id
        )
        db.session.add(new_comment)
        db.session.commit()

        return jsonify({'status': 'success', 'message': 'Komentar berhasil disimpan.'}), 200

    except Exception as e:
        db.session.rollback()
        print('Error saat menyimpan komentar:', e)
        return jsonify({'status': 'error', 'message': f"Gagal menyimpan komentar: {str(e)}"}), 500
    
@app.route('/api/save_row_action', methods=['POST'])
@login_required
def api_save_row_action():
    try:
        # --- Langkah 1: Ambil dan Validasi Data JSON ---
        data = request.get_json()

        # --- TAMBAHAN: Pengecekan Tipe Data untuk Mencegah Error ---
        if not isinstance(data, dict):
            return jsonify({"error": "Format data yang dikirim tidak valid. Harus berupa JSON objek."}), 400

        # --- Ambil Data dari Payload ---
        folder_name = data.get('folder_name')
        filename = data.get('filename')
        owner_id = data.get('owner_id')
        row_id = data.get('row_id')
        is_ganti = data.get('is_ganti', False)
        pic_user_id = data.get('pic_user_id')

        # --- Validasi Kelengkapan Data ---
        if not all([folder_name, filename, owner_id is not None, row_id is not None]):
            return jsonify({"error": "Data tidak lengkap. Diperlukan: folder_name, filename, owner_id, row_id."}), 400

        # --- Cek Izin Akses ---
        is_owner = (str(current_user.id) == str(owner_id))
        is_shared_to_me = SharedFolder.query.filter_by(
            owner_id=owner_id, 
            folder_name=folder_name, 
            shared_with_id=current_user.id
        ).first()

        if not is_owner and not is_shared_to_me:
            return jsonify({"error": "Akses ditolak. Anda tidak memiliki izin untuk file ini."}), 403

        # --- Operasi Database: Cari atau Buat Data Baru ---
        action = RowAction.query.filter_by(
            owner_id=owner_id,
            folder_name=folder_name,
            filename=filename,
            row_id=row_id
        ).first()

        if action:
            # Update data yang sudah ada
            action.is_ganti = is_ganti
            action.pic_user_id = pic_user_id
        else:
            # Buat data baru
            action = RowAction(
                owner_id=owner_id,
                folder_name=folder_name,
                filename=filename,
                row_id=row_id,
                is_ganti=is_ganti,
                pic_user_id=pic_user_id
            )
            db.session.add(action)
        
        db.session.commit()
        
        return jsonify({"status": "success", "message": "Status baris berhasil disimpan."}), 200

    except Exception as e:
        db.session.rollback()
        # --- Cetak error lengkap ke terminal untuk debugging ---
        import traceback
        print("!!! ERROR TERJADI DI api_save_row_action !!!")
        traceback.print_exc() # Ini akan mencetak error lengkap ke terminal
        # --- Kirim pesan error yang jelas ke frontend ---
        return jsonify({"error": f"Terjadi kesalahan di server: {str(e)}"}), 500

    except Exception as e:
        db.session.rollback()
        # --- Cetak error lengkap ke terminal untuk debugging ---
        import traceback
        print("!!! ERROR TERJADI DI api_save_row_action !!!")
        traceback.print_exc()
        # --- Kirim pesan error yang jelas ke frontend ---
        return jsonify({"error": f"Terjadi kesalahan di server: {str(e)}"}), 500

    except Exception as e:
        db.session.rollback()
        print(f"Error saat menyimpan status baris: {e}")
        return jsonify({"error": f"Gagal menyimpan status: {str(e)}"}), 500

@app.route('/log_analysis')
@login_required
def log_analysis_page():
    """Menampilkan halaman log analisis dan tracker tugas."""
    return render_template('log_analysis.html', username=current_user.username, label=current_user.label)

app.route('/api/log_analysis_start', methods=['POST'])
@login_required
def api_log_analysis_start():
    """Mencatat awal proses analisis."""
    data = request.json
    filename = data.get('filename')
    feature_type = data.get('feature_type')

    if not filename or not feature_type:
        return jsonify({"error": "Data tidak lengkap."}), 400

    try:
        new_log = AnalysisLog(
            user_id=current_user.id,
            filename=filename,
            feature_type=feature_type,
            status='unfinished'
        )
        db.session.add(new_log)
        db.session.commit()
        
        # Kembalikan ID log untuk digunakan di frontend
        return jsonify({
            "status": "success",
            "log_id": new_log.id
        }), 201
    except Exception as e:
        db.session.rollback()
        print(f"Error saat memulai log analisis: {e}")
        return jsonify({"error": "Gagal memulai log."}), 500

@app.route('/api/log_analysis_end', methods=['POST'])
@login_required
def api_log_analysis_end():
    """Mencatat akhir proses analisis."""
    data = request.json
    log_id = data.get('log_id')
    status = data.get('status') # 'done' atau 'error'

    if not log_id or status not in ['done', 'error']:
        return jsonify({"error": "Data tidak lengkap atau status tidak valid."}), 400

    try:
        log_entry = AnalysisLog.query.get(log_id)
        if not log_entry or log_entry.user_id != current_user.id:
            return jsonify({"error": "Log tidak ditemukan atau akses ditolak."}), 404
        
        log_entry.end_time = datetime.datetime.utcnow()
        log_entry.status = status
        db.session.commit()
        
        return jsonify({"status": "success"}), 200
    except Exception as e:
        db.session.rollback()
        print(f"Error saat mengakhiri log analisis: {e}")
        return jsonify({"error": "Gagal mengakhiri log."}), 500

@app.route('/api/add_manual_task', methods=['POST'])
@login_required
def api_add_manual_task():
    """Menambahkan tugas manual ke log."""
    data = request.get_json()
    if not data:
        return jsonify({"error": "Request body tidak valid (harus berupa JSON)."}), 400

    filename = data.get('filename')
    feature_type = data.get('feature_type')
    start_date_str = data.get('start_time')
    deadline_str = data.get('deadline')
    end_date_str = data.get('end_time')

    if not filename:
        return jsonify({"error": "Nama Dokumen harus diisi."}), 400
    
    if not start_date_str:
        return jsonify({"error": "Tanggal Mulai harus diisi."}), 400

    try:
        start_time = parse_flexible_date(start_date_str)
        deadline_date = parse_flexible_date(deadline_str)
        deadline = deadline_date.date() if deadline_date else None
        end_time = parse_flexible_date(end_date_str)
        
        # --- GUNAKAN FUNGSI HELPER UNTUK MENENTUKAN STATUS ---
        status = calculate_task_status(start_time, deadline, end_time)
        
        new_log = AnalysisLog(
            user_id=current_user.id,
            document_type=data.get('document_type'),
            filename=filename,
            feature_type=feature_type,
            start_time=start_time,
            deadline=deadline,
            end_time=end_time,
            status=status
        )
        db.session.add(new_log)
        db.session.commit()
        
        return jsonify({"status": "success", "message": "Tugas berhasil ditambahkan."}), 201

    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        db.session.rollback()
        print(f"Error saat menambah tugas manual: {e}")
        return jsonify({"error": "Gagal menambah tugas."}), 500   

@app.route('/api/get_analysis_logs', methods=['GET'])
@login_required
def api_get_analysis_logs():
    """Mengambil semua log analisis untuk user saat ini."""
    try:
        logs = AnalysisLog.query.filter_by(user_id=current_user.id).order_by(AnalysisLog.start_time.desc()).all()
        
        logs_data = []
        for log in logs:
            # --- PERBAIKAN: Tambahkan data deadline ---
            deadline_str = ''
            if log.deadline:
                deadline_str = log.deadline.strftime('%d %b %Y, %H:%M')
            
            logs_data.append({
                "id": log.id,
                "filename": log.filename,
                "feature_type": log.feature_type,
                "start_time": log.start_time.strftime('%d %b %Y, %H:%M'),
                "deadline": deadline_str, # <<<< KIRIMKAN DATA DEADLINE
                "end_time": log.end_time.strftime('%d %b %Y, %H:%M') if log.end_time else None,
                "status": log.status
            })
        return jsonify(logs_data), 200
    except Exception as e:
        print(f"Error saat mengambil log analisis: {e}")
        return jsonify({"error": "Gagal mengambil log."}), 500

@app.route('/api/edit_task/<int:log_id>', methods=['POST'])
@login_required
def api_edit_task(log_id):
    """Mengedit tugas yang sudah ada."""
    log = AnalysisLog.query.get(log_id)
    if not log or log.user_id != current_user.id:
        return jsonify({"error": "Tugas tidak ditemukan atau akses ditolak."}), 404

    data = request.get_json()
    if not data:
        return jsonify({"error": "Request body tidak valid (harus berupa JSON)."}), 400

    # Ambil data dari form, gunakan nilai lama jika tidak ada di form
    log.filename = data.get('filename', log.filename)
    log.feature_type = data.get('feature_type', log.feature_type)
    log.document_type = data.get('document_type', log.document_type)
    
    try:
        # Parse tanggal yang dikirim dari form
        start_time = parse_flexible_date(data.get('start_time'))
        deadline_date = parse_flexible_date(data.get('deadline'))
        end_time = parse_flexible_date(data.get('end_time'))

        # Update data di objek log
        if start_time:
            log.start_time = start_time
        
        if deadline_date:
            log.deadline = deadline_date.date()
        else:
            log.deadline = None
            
        if end_time:
            log.end_time = end_time
        else:
            log.end_time = None

        log.status = calculate_task_status(log.start_time, log.deadline, log.end_time)
        
        db.session.commit()
        return jsonify({"status": "success", "message": "Tugas berhasil diperbarui."}), 200
    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        db.session.rollback()
        print(f"Error saat mengedit tugas: {e}")
        return jsonify({"error": f"Gagal memperbarui tugas: {e}"}), 500

@app.route('/api/delete_task/<int:log_id>', methods=['DELETE'])
@login_required
def api_delete_task(log_id):
    log = AnalysisLog.query.get(log_id)
    if not log or log.user_id != current_user.id:
        return jsonify({"error": "Tugas tidak ditemukan atau akses ditolak."}), 404

    try:
        db.session.delete(log)
        db.session.commit()
        return jsonify({"status": "success", "message": "Tugas berhasil dihapus."}), 200
    except Exception as e:
        db.session.rollback()
        print(f"Error saat menghapus tugas: {e}")
        return jsonify({"error": "Gagal menghapus tugas."}), 500

@app.route('/mailbox')
@login_required
def mailbox_page():
    """Menampilkan halaman Mailbox."""
    return render_template('mailbox.html', 
                           username=current_user.username, 
                           label=current_user.label)

@app.route('/api/send_message', methods=['POST'])
@login_required
def api_send_message():
    """Mengirim pesan ke user lain."""
    recipient_id = request.form.get('recipient_id')
    subject = request.form.get('subject')
    body = request.form.get('body')
    attachment = request.files.get('attachment')

    if not recipient_id or not subject:
        return jsonify({"error": "Penerima dan Subjek harus diisi."}), 400

    try:
        attachment_path = None
        original_filename = None
        if attachment and attachment.filename:
            original_filename = attachment.filename
            # Buat nama file unik untuk menghindari tabrakan
            filename = f"{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}_{current_user.id}_{original_filename}"
            attachment_path = os.path.join('mailbox_attachments', filename)
            attachment.save(attachment_path)

        new_message = Message(
            sender_id=current_user.id,
            recipient_id=recipient_id,
            subject=subject,
            body=body,
            attachment_path=attachment_path,
            original_filename=original_filename
        )
        db.session.add(new_message)
        db.session.commit()

        return jsonify({"status": "success", "message": "Pesan berhasil dikirim."}), 201

    except Exception as e:
        db.session.rollback()
        print(f"Error saat mengirim pesan: {e}")
        return jsonify({"error": "Gagal mengirim pesan."}), 500

@app.route('/api/delete_message/<int:message_id>', methods=['DELETE'])
@login_required
def api_delete_message(message_id):
    """Menghapus pesan (hanya pengirim yang bisa hapus)."""
    message = Message.query.get(message_id)
    if not message or message.sender_id != current_user.id:
        return jsonify({"error": "Pesan tidak ditemukan atau Anda tidak memiliki izin."}), 404

    try:
        # Jika ada lampiran, hapus file-nya dari server
        if message.attachment_path and os.path.exists(message.attachment_path):
            os.remove(message.attachment_path)
        
        db.session.delete(message)
        db.session.commit()
        return jsonify({"status": "success", "message": "Pesan berhasil dihapus."}), 200
    except Exception as e:
        db.session.rollback()
        print(f"Error saat menghapus pesan: {e}")
        return jsonify({"error": "Gagal menghapus pesan."}), 500

@app.route('/api/get_messages', methods=['POST'])
@login_required
def api_get_messages():
    """Mengambil daftar pesan (inbox atau sent)."""
    data = request.json
    msg_type = data.get('type') # 'inbox' atau 'sent'

    if msg_type == 'inbox':
        messages = Message.query.filter_by(recipient_id=current_user.id).order_by(Message.timestamp.desc()).all()
    elif msg_type == 'sent':
        messages = Message.query.filter_by(sender_id=current_user.id).order_by(Message.timestamp.desc()).all()
    else:
        return jsonify({"error": "Tipe pesan tidak valid."}), 400

    messages_data = []
    for msg in messages:
        messages_data.append({
            "id": msg.id,
            "other_user": msg.sender.username if msg_type == 'inbox' else msg.recipient.username,
            "subject": msg.subject,
            "body": msg.body,
            "timestamp": msg.timestamp.isoformat(), 
            "is_read": msg.is_read,
            "has_attachment": bool(msg.original_filename)
        })
    return jsonify(messages_data), 200

@app.route('/api/download_message_attachment/<int:message_id>')
@login_required
def api_download_message_attachment(message_id):
    """Mengunduh lampiran pesan."""
    message = Message.query.get(message_id)
    if not message or (message.sender_id != current_user.id and message.recipient_id != current_user.id):
        return jsonify({"error": "Akses ditolak."}), 403

    if not message.attachment_path or not os.path.exists(message.attachment_path):
        return jsonify({"error": "File tidak ditemukan."}), 404

    return send_file(message.attachment_path, as_attachment=True, download_name=message.original_filename)

@app.route('/api/get_unread_count', methods=['GET'])
@login_required
def api_get_unread_count():
    """Menghitung jumlah pesan yang belum dibaca untuk user saat ini."""
    try:
        count = Message.query.filter_by(recipient_id=current_user.id, is_read=False).count()
        return jsonify({"count": count}), 200
    except Exception as e:
        print(f"Error saat menghitung pesan belum dibaca: {e}")
        return jsonify({"error": "Gagal menghitung pesan."}), 500

@app.route('/api/mark_message_read/<int:message_id>', methods=['POST'])
@login_required
def api_mark_message_read(message_id):
    """Menandai pesan sebagai sudah dibaca."""
    message = Message.query.get(message_id)
    if not message or message.recipient_id != current_user.id:
        return jsonify({"error": "Pesan tidak ditemukan atau akses ditolak."}), 404

    try:
        message.is_read = True
        db.session.commit()
        return jsonify({"status": "success"}), 200
    except Exception as e:
        db.session.rollback()
        print(f"Error saat menandai pesan sebagai dibaca: {e}")
        return jsonify({"error": "Gagal memperbarui status pesan."}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)